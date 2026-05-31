import os
import re
from collections import defaultdict

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    MessageHandler,
    ContextTypes,
    filters,
)

# =====================================================
# TOKEN
# =====================================================

BOT_TOKEN = "8543306095:AAEBpo1SKrQJWqtyw_ik9rRcX7phkRj3B1g"

# =====================================================
# ИСКЛЮЧЕНИЯ
# =====================================================

IGNORE_WORDS = [
    "ЙОГУРТ",
    "НАПИТОК",
    "ЧАЙ",
    "ХЛЕБ",
    "САХАР",
    "КЕФИР",
    "ПЕЧЕНЬЕ",
    "КОМПОТ",
    "СОК",
    "ВОДА",
    "МОРС",
    "КИСЕЛЬ",
    "ЛИМОН",
    # "МАСЛО" убрано — ловит «РИС ОТВАРНОЙ С ЗЕЛЕНЬЮ И СЛИВОЧНЫМ МАСЛОМ»
    "БУЛОЧКА",
    "ЯЙЦО",
    "ЯБЛОКО",
    "ФРУКТЫ",
]

IGNORE_SECTIONS = [
    "ЗАВТРАК",
    "ЗАВТРАК 2",
    "ПОЛДНИК",
    "2-Й УЖИН",
]

SOUP_KEYWORDS = [
    "СУП",
    "БОРЩ",
    "БУЛЬОН",
    "КРЕМ",
    "УХА",
    "ЩАВЕЛ",
    "РАССОЛЬНИК",
    "СОЛЯНКА",
]

COLD_KEYWORDS = [
    "САЛАТ",
    "МИМОЗА",
    "МИКС",
    "СВЕКОЛЬНЫЙ",
    # "ЗАПЕЧЕНН"/"ОВОЩИ"/"ОВОЩНОЙ" убраны — ловят горячие блюда ужина
    "КАПУСТ",
    "ЗАКУСКА",
    "НАРЕЗКА",
]

GARNISH_KEYWORDS = [
    "РИС",
    "ПЮРЕ",
    "ГРЕЧ",
    "КАРТОФ",
    "БУЛГУР",
    "КИНОА",
    "КУСКУС",
    "КУС КУС",
    "ПАСТА",
    "СПАГЕТТИ",
    "МАКАРОН",
    "ЧЕЧЕВИЦ",
    "ФАСОЛ",
    "ПЕРЛОВК",
    "ЯЧНЕВАЯ",
]

EXCLUDE_DISHES = [
    "БОУЛ",
]

# =====================================================
# ВСПОМОГАТЕЛЬНЫЕ
# =====================================================

def normalize(text):
    if not text:
        return ""
    return str(text).upper().strip()


def should_ignore(dish):
    dish_up = normalize(dish)
    return any(w in dish_up for w in IGNORE_WORDS)


def should_exclude(dish):
    dish_up = normalize(dish)
    return any(w in dish_up for w in EXCLUDE_DISHES)


def is_soup(dish):
    dish_up = normalize(dish)
    if any(w in dish_up for w in SOUP_KEYWORDS):
        return True
    if dish_up.startswith("ЩИ"):
        return True
    return False


def is_cold(dish):
    dish_up = normalize(dish)
    return any(w in dish_up for w in COLD_KEYWORDS)


def is_garnish(dish):
    dish_up = normalize(dish)
    # Not a soup, not cold — then check garnish keywords
    return (
        any(w in dish_up for w in GARNISH_KEYWORDS)
        and not is_soup(dish)
        and not is_cold(dish)
    )


def extract_date(filename):
    match = re.search(r"(\d{1,2}[._]\d{1,2}([._]\d{2,4})?)", filename)
    if match:
        raw = match.group(1).replace("_", ".")
        return raw
    return "дата"

# =====================================================
# ЯЧЕЙКИ ТАБЛИЦЫ — вспомогательные
# =====================================================

def set_cell_bold(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    # Если runs нет — создаём
    if not cell.paragraphs[0].runs:
        run = cell.paragraphs[0].add_run(cell.text)
        run.bold = True
        cell.paragraphs[0].runs[0].text = ""
        cell.paragraphs[0].add_run(cell.text).bold = True


def set_bold_text(cell, text):
    """Очищаем ячейку и вставляем жирный текст."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True


def add_header_row(table, date_str):
    """Первая строка таблицы: дата | КРАФТ | ПЛАСТИК."""
    row = table.rows[0].cells
    set_bold_text(row[0], date_str)
    set_bold_text(row[1], "КРАФТ")
    set_bold_text(row[2], "ПЛАСТИК")


def add_section_row(table, name):
    """Строка-заголовок секции: объединённые ячейки, текст по центру, без внутренних разделителей."""
    row = table.add_row()
    cells = row.cells

    # Объединяем все три ячейки
    cells[0].merge(cells[2])

    # Текст по центру, жирный
    para = cells[0].paragraphs[0]
    para.clear()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # Убираем левый и правый внутренние разделители (оставляем только внешние top/bottom)
    tc = cells[0]._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    no_border = {"w:val": "nil"}
    for side in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_dish_row(table, dish, kraft, plastic):
    """Строка блюда."""
    row = table.add_row().cells
    row[0].text = dish
    row[1].text = str(int(kraft)) if kraft and kraft > 0 else ""
    row[2].text = str(int(plastic)) if plastic and plastic > 0 else ""

# =====================================================
# ОСНОВНАЯ ЛОГИКА — ПАРСИНГ EXCEL
# =====================================================

def parse_excel(excel_path):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    ws = wb.active; mc = ws.max_column or 30

    packaging_row = 1
    for r in range(1, 10):
        text = "".join(str(ws.cell(r,c).value or "").upper() for c in range(1, mc+1))
        if "КРАФТ" in text or "ПЛАСТИК" in text:
            packaging_row = r; break

    plastic_cols, kraft_cols = set(), set()
    for c in range(1, mc+1):
        hdr = normalize(ws.cell(packaging_row, c).value)
        if "ПЛАСТИК" in hdr: plastic_cols.add(c)
        elif "КРАФТ" in hdr: kraft_cols.add(c)

    patient_row = packaging_row + 1
    plastic_osnov_cols = set()
    for c in plastic_cols:
        val = normalize(ws.cell(patient_row, c).value)
        if "ОСНОВНОЙ" in val: plastic_osnov_cols.add(c)

    SKIP = set(IGNORE_SECTIONS)
    current_section = None
    section_data = defaultdict(list)  # (dish, col_qty, row_type, has_plastic_osnov)

    for r in range(1, ws.max_row+1):
        row_vals = [ws.cell(r,c).value for c in range(1, mc+1)]
        first  = normalize(row_vals[0]) if row_vals[0] else ""
        second = normalize(row_vals[1]) if len(row_vals)>1 and row_vals[1] else ""
        found = False
        for sec in ["2-Й УЖИН","ЗАВТРАК 2","ЗАВТРАК","ОБЕД","ПОЛДНИК","УЖИН","ХОЛОДНЫЙ"]:
            if sec in first or sec in second:
                current_section = sec; found = True; break
        if found: continue
        if current_section is None or current_section in SKIP: continue

        dish = ""
        if row_vals[1]: dish = str(row_vals[1]).strip()
        elif len(row_vals)>2 and row_vals[2]: dish = str(row_vals[2]).strip()
        dish = normalize(dish)
        if not dish or dish == "NONE": continue
        if should_ignore(dish): continue
        if any(x in dish for x in [normalize(s) for s in IGNORE_SECTIONS]): continue
        if should_exclude(dish): continue

        col_qty = {}; has_plastic_osnov = False
        for ci, val in enumerate(row_vals):
            col = ci + 1
            if not isinstance(val, (int,float)) or val <= 0: continue
            qty = int(val)
            if col in plastic_cols:
                col_qty[col] = col_qty.get(col,0) + qty
                if col in plastic_osnov_cols: has_plastic_osnov = True
            elif col in kraft_cols:
                col_qty[col] = col_qty.get(col,0) + qty
        if not col_qty: continue

        if is_soup(dish):                                          row_type = "soup"
        elif current_section == "ХОЛОДНЫЙ":                       row_type = "cold"
        elif is_cold(dish):                                        row_type = "cold"
        elif dish.startswith("ОВОЩИ") and current_section=="ОБЕД": row_type = "cold"
        elif dish.startswith("ОВОЩИ") and current_section=="УЖИН": row_type = "garnish"
        elif is_garnish(dish):                                     row_type = "garnish"
        else:                                                      row_type = "main"

        section_data[current_section].append((dish, col_qty, row_type, has_plastic_osnov))

    return section_data, plastic_cols, kraft_cols, plastic_osnov_cols

# =====================================================
# СОПОСТАВЛЕНИЕ ГОРЯЧЕГО С ГАРНИРОМ
# =====================================================

def pair_by_column(rows, plastic_cols, kraft_cols, plastic_osnov_cols):
    """Сопоставляет блюдо с гарниром по колонке пациента, а не позиционно."""
    soup_rows    = [(d,cq)    for d,cq,t,h in rows if t=="soup"]
    main_rows    = [(d,cq)    for d,cq,t,h in rows if t=="main"]
    garnish_rows = [(d,cq)    for d,cq,t,h in rows if t=="garnish"]
    cold_rows    = [(d,cq,h)  for d,cq,t,h in rows if t=="cold"]

    soup_agg = defaultdict(lambda:[0,0])
    for d,cq in soup_rows:
        for qty in cq.values(): soup_agg[d][1] += qty
    soups = [(d,k,p) for d,(k,p) in soup_agg.items()]

    cold_agg = defaultdict(lambda:[0,0,False])
    for d,cq,h in cold_rows:
        for col,qty in cq.items():
            if col in plastic_cols:
                cold_agg[d][1]+=qty
                if col in plastic_osnov_cols: cold_agg[d][2]=True
            else: cold_agg[d][0]+=qty
        if h: cold_agg[d][2]=True
    cold = [(d,k,p,h) for d,(k,p,h) in cold_agg.items()]

    col_to_main    = {}
    col_to_garnish = {}
    for d,cq in main_rows:
        for col,qty in cq.items(): col_to_main[col] = (d, qty)
    for d,cq in garnish_rows:
        for col in cq: col_to_garnish[col] = d

    pair_agg = defaultdict(lambda:[0,0])
    for col,(md,qty) in col_to_main.items():
        g = col_to_garnish.get(col)
        key = (md, g)
        if col in plastic_cols: pair_agg[key][1] += qty
        else:                   pair_agg[key][0] += qty

    paired = [(f"{m} и {g}" if g else m, k, p) for (m,g),(k,p) in pair_agg.items()]
    return soups, paired, cold

# =====================================================
# ГЕНЕРАЦИЯ WORD
# =====================================================

def generate_word(excel_path):
    section_data, plastic_cols, kraft_cols, plastic_osnov_cols = parse_excel(excel_path)

    lunch_rows = section_data.get("ОБЕД", [])
    lunch_soups, lunch_hot, lunch_cold = pair_by_column(
        lunch_rows, plastic_cols, kraft_cols, plastic_osnov_cols
    )
    dinner_rows = section_data.get("УЖИН", [])
    dinner_soups, dinner_hot, dinner_cold = pair_by_column(
        dinner_rows, plastic_cols, kraft_cols, plastic_osnov_cols
    )
    холодный_rows = section_data.get("ХОЛОДНЫЙ", [])
    cold_dishes = []
    for d, k, p, h in lunch_cold:
        if not h: cold_dishes.append((d, k + p))
    for d, k, p, h in dinner_cold:
        if not h: cold_dishes.append((d, k + p))
    for d, cq, t, h in холодный_rows:
        if not h: cold_dishes.append((d, sum(cq.values())))

    # --------------------------------------------------
    # Создаём документ — одна большая таблица
    # --------------------------------------------------
    doc = Document()

    # Убираем отступы страницы для компактности
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    date_str = extract_date(os.path.basename(excel_path))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Ширины колонок (в twips): общая ≈ 9200 при полях 2 см
    col_widths_cm = [13, 2.5, 2.5]
    for i, cell in enumerate(table.rows[0].cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(col_widths_cm[i] * 567)))  # 1 см ≈ 567 twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    # Шапка
    add_header_row(table, date_str)

    # --- ОБЕД ---
    add_section_row(table, "обед")
    for dish, k, p in lunch_soups:
        add_dish_row(table, dish, 0, k + p)
    for dish, k, p in lunch_hot:
        add_dish_row(table, dish, k, p)
    for _ in range(4):
        add_dish_row(table, "", 0, 0)

    # --- УЖИН ---
    add_section_row(table, "ужин")
    for dish, k, p in dinner_soups:
        add_dish_row(table, dish, 0, k + p)
    for dish, k, p in dinner_hot:
        add_dish_row(table, dish, k, p)
    for _ in range(4):
        add_dish_row(table, "", 0, 0)

    # --- ХОЛОДНЫЙ ЦЕХ ---
    add_section_row(table, "Холодный цех")
    for dish, qty in cold_dishes:
        add_dish_row(table, dish, 0, qty)
    for _ in range(4):
        add_dish_row(table, "", 0, 0)

    # Шрифт, размер и отступы для всей таблицы
    for row in table.rows:
        for cell in row.cells:
            # Внутренние отступы ячейки
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side, val in [("top", "60"), ("bottom", "60"), ("left", "120"), ("right", "120")]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), val)
                el.set(qn("w:type"), "dxa")
                tcMar.append(el)
            tcPr.append(tcMar)

            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Arial"
                # Если параграф без runs (пустой), всё равно задаём шрифт через XML
                if not para.runs:
                    pPr = para._p.get_or_add_pPr()
                    rPr = OxmlElement("w:rPr")
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), "Arial")
                    rFonts.set(qn("w:hAnsi"), "Arial")
                    rPr.append(rFonts)
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), "22")
                    rPr.append(sz)
                    pPr.append(rPr)

    output_path = f"лист фасовки {date_str}.docx"
    doc.save(output_path)
    return output_path

# =====================================================
# TELEGRAM
# =====================================================

async def handle_file(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
):
    document = update.message.document

    if not document.file_name.lower().endswith((".xlsx", ".xls")):
        await update.message.reply_text("Отправьте Excel файл (.xlsx или .xls)")
        return

    await update.message.reply_text("Файл получен. Формирую лист фасовки...")

    file = await document.get_file()
    input_path = "input.xlsx"
    await file.download_to_drive(input_path)

    try:
        output_path = generate_word(input_path)
        await update.message.reply_document(
            document=open(output_path, "rb"),
            caption=f"✅ Готово: {os.path.basename(output_path)}",
        )
    except Exception as e:
        await update.message.reply_text(f"❌ Ошибка:\n{str(e)}")

# =====================================================
# ЗАПУСК
# =====================================================

def main():
    app = ApplicationBuilder().token(BOT_TOKEN).build()
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    print("🤖 БОТ ЗАПУЩЕН")
    app.run_polling()


if __name__ == "__main__":
    main()